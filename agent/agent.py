#!/usr/bin/env python3
"""CursorEye Local Screen Agent — runs on user's Mac.
Connects to CursorEye cloud via WebSocket, executes screen commands.
All screen operations use native macOS APIs."""

import asyncio
import base64
import json
import os
import subprocess
import sys
import tempfile
import websockets

AGENT_VERSION = "2.1.0"
DEFAULT_WS_URL = "ws://localhost:8765"
WORKER_WINDOW_TITLE = "🔴 CursorEye Worker"
PURPLE_CURSOR_PID_FILE = os.path.expanduser("~/.cursoreye/purple_cursor.pid")
PURPLE_CURSOR_POS_FILE = os.path.expanduser("~/.cursoreye/purple_cursor_pos.json")

# ─── Screenshot ───────────────────────────────────────────────────────────


def take_screenshot(region=None):
    """Take screenshot, return base64 PNG."""
    cmd = ["screencapture", "-x", "-t", "png"]
    if region:
        cmd += [
            "-R",
            str(region["x"]),
            str(region["y"]),
            str(region["w"]),
            str(region["h"]),
        ]
    tmp = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
    tmp.close()
    cmd.append(tmp.name)
    subprocess.run(cmd, check=True)
    with open(tmp.name, "rb") as f:
        data = base64.b64encode(f.read()).decode()
    os.unlink(tmp.name)
    return data


# ─── OCR via Apple Vision Framework ────────────────────────────────────────


def ocr_screen(region=None):
    """Screenshot + OCR, return extracted text with bounding boxes."""
    img_b64 = take_screenshot(region)
    return _ocr_from_base64(img_b64)


def _ocr_from_base64(img_b64):
    """Run OCR on a base64 PNG image."""
    swift_script = (
        """
import Vision
import AppKit
import Foundation

guard let url = URL(string: "data:image/png;base64,%s"),
    let data = try? Data(contentsOf url),
    let image = NSImage(data: data),
    let cgImage = image.cgImage(forProposedRect: nil, context: nil, hints: nil)
else { print("[]"); exit(0) }

let request = VNRecognizeTextRequest { req, error in
    guard let observations = req.results as? [VNRecognizedTextObservation] else { print("[]"); return }
    var results = [[String: Any]]()
    for obs in observations {
        guard let text = obs.topCandidates(1).first?.string else { continue }
        let bbox = obs.boundingBox
        results.append([
            "text": text,
            "x": Int(bbox.origin.x * 100),
            "y": Int((1 - bbox.origin.y - bbox.height) * 100),
            "w": Int(bbox.width * 100),
            "h": Int(bbox.height * 100)
        ])
    }
    let jsonData = try! JSONSerialization.data(withJSONObject: results)
    print(String(data: jsonData, encoding: .utf8)!)
}

request.recognitionLevel = .accurate
request.usesLanguageCorrection = true

let handler = VNImageRequestHandler(cgImage: cgImage, options: [:])
try? handler.perform([request])
"""
        % img_b64
    )

    result = subprocess.run(
        ["swift", "-"], input=swift_script, capture_output=True, text=True, timeout=30
    )
    try:
        return {
            "text_blocks": json.loads(result.stdout.strip())
            if result.stdout.strip()
            else [],
            "image": img_b64,
        }
    except json.JSONDecodeError:
        return {"text_blocks": [], "image": img_b64, "raw_ocr": result.stdout[:500]}


# ─── Worker Window Capture (AI's own "computer") ──────────────────────────


def get_worker_window_bounds():
    """Get the bounds of the CursorEye Worker terminal window via AppleScript."""
    script = f"""
    tell application "System Events"
        set p to first process whose name is "Terminal"
        set w to first window of p whose name contains "CursorEye Worker"
        set pos to position of w
        set sz to size of w
        return (item 1 of pos) & "," & (item 2 of pos) & "," & (item 1 of sz) & "," & (item 2 of sz)
    end tell
    """
    result = subprocess.run(["osascript", "-e", script], capture_output=True, text=True)
    if result.returncode != 0:
        return None
    try:
        parts = result.stdout.strip().split(", ")
        return {
            "x": int(parts[0]),
            "y": int(parts[1]),
            "w": int(parts[2]),
            "h": int(parts[3]),
        }
    except (ValueError, IndexError):
        return None


def screenshot_worker():
    """Screenshot the Worker window only — AI's own screen."""
    bounds = get_worker_window_bounds()
    if not bounds:
        return {
            "success": False,
            "error": "Worker window not found. Open it first with open_worker.",
        }
    img_b64 = take_screenshot(bounds)
    return {
        "success": True,
        "action": "screenshot_worker",
        "image": img_b64[:100] + "...",
        "image_length": len(img_b64),
        "window_bounds": bounds,
    }


def ocr_worker():
    """OCR the Worker window only — AI can read its own workspace."""
    bounds = get_worker_window_bounds()
    if not bounds:
        return {
            "success": False,
            "error": "Worker window not found. Open it first with open_worker.",
        }
    img_b64 = take_screenshot(bounds)
    ocr_result = _ocr_from_base64(img_b64)
    texts = [b["text"] for b in ocr_result.get("text_blocks", [])]
    return {
        "success": True,
        "action": "ocr_worker",
        "text": "\n".join(texts),
        "blocks": len(texts),
    }


# ─── Mouse Control via Quartz ──────────────────────────────────────────────


def mouse_click(x_pct, y_pct, button="left", click_count=1):
    """Click at percentage position on screen."""
    import Quartz

    screens = Quartz.CGDisplayBounds(Quartz.CGMainDisplayID())
    x = int(screens.size.width * x_pct / 100)
    y = int(screens.size.height * y_pct / 100)
    btn_map = {
        "left": Quartz.kCGMouseButtonLeft,
        "right": Quartz.kCGMouseButtonRight,
        "center": Quartz.kCGMouseButtonCenter,
    }
    btn = btn_map.get(button, Quartz.kCGMouseButtonLeft)
    for _ in range(click_count):
        Quartz.CGEventCreateMouseEvent(None, Quartz.kCGEventMouseMoved, (x, y), btn)
        Quartz.CGEventPost(
            Quartz.kCGHIDEventTap,
            Quartz.CGEventCreateMouseEvent(
                None, Quartz.kCGEventLeftMouseDown, (x, y), btn
            ),
        )
        Quartz.CGEventPost(
            Quartz.kCGHIDEventTap,
            Quartz.CGEventCreateMouseEvent(
                None, Quartz.kCGEventLeftMouseUp, (x, y), btn
            ),
        )
    return {"success": True, "x": x, "y": y}


def mouse_drag(x1_pct, y1_pct, x2_pct, y2_pct, duration=0.5):
    """Drag from one position to another."""
    import Quartz

    screens = Quartz.CGDisplayBounds(Quartz.CGMainDisplayID())
    x1 = int(screens.size.width * x1_pct / 100)
    y1 = int(screens.size.height * y1_pct / 100)
    x2 = int(screens.size.width * x2_pct / 100)
    y2 = int(screens.size.height * y2_pct / 100)
    Quartz.CGEventPost(
        Quartz.kCGHIDEventTap,
        Quartz.CGEventCreateMouseEvent(
            None, Quartz.kCGEventLeftMouseDown, (x1, y1), Quartz.kCGMouseButtonLeft
        ),
    )
    steps = 20
    import time

    for i in range(1, steps + 1):
        t = i / steps
        mx = int(x1 + (x2 - x1) * t)
        my = int(y1 + (y2 - y1) * t)
        Quartz.CGEventPost(
            Quartz.kCGHIDEventTap,
            Quartz.CGEventCreateMouseEvent(
                None,
                Quartz.kCGEventLeftMouseDragged,
                (mx, my),
                Quartz.kCGMouseButtonLeft,
            ),
        )
        time.sleep(duration / steps)
    Quartz.CGEventPost(
        Quartz.kCGHIDEventTap,
        Quartz.CGEventCreateMouseEvent(
            None, Quartz.kCGEventLeftMouseUp, (x2, y2), Quartz.kCGMouseButtonLeft
        ),
    )
    return {"success": True}


def mouse_scroll(x_pct, y_pct, amount, direction="down"):
    """Scroll at position."""
    import Quartz

    screens = Quartz.CGDisplayBounds(Quartz.CGMainDisplayID())
    x = int(screens.size.width * x_pct / 100)
    y = int(screens.size.height * y_pct / 100)
    scroll_val = -amount if direction == "down" else amount
    for _ in range(abs(amount)):
        event = Quartz.CGEventCreateScrollWheelEvent(
            None, Quartz.kCGScrollEventUnitLine, 1, scroll_val
        )
        Quartz.CGEventPost(Quartz.kCGHIDEventTap, event)
    return {"success": True}


# ─── Keyboard Control ──────────────────────────────────────────────────────


def keyboard_type(text):
    """Type text character by character."""
    import Quartz

    for char in text:
        keycode_map = {
            " ": 49,
            "\n": 36,
            "\t": 48,
            ".": 47,
            ",": 43,
            "/": 44,
            ";": 41,
            "'": 39,
            "[": 33,
            "]": 30,
            "\\": 42,
            "-": 27,
            "=": 24,
            "`": 50,
            "1": 18,
            "2": 19,
            "3": 20,
            "4": 21,
            "5": 23,
            "6": 22,
            "7": 26,
            "8": 28,
            "9": 25,
            "0": 29,
        }
        if char in keycode_map:
            key_code = keycode_map[char]
            shift = False
        elif char.isupper():
            key_code = keycode_map.get(char.lower(), 0)
            shift = True
        else:
            key_code = keycode_map.get(char, 0)
            shift = False
        if key_code == 0:
            continue
        if shift:
            shift_event = Quartz.CGEventCreateKeyboardEvent(None, 56, True)
            Quartz.CGEventPost(Quartz.kCGHIDEventTap, shift_event)
        down = Quartz.CGEventCreateKeyboardEvent(None, key_code, True)
        Quartz.CGEventPost(Quartz.kCGHIDEventTap, down)
        up = Quartz.CGEventCreateKeyboardEvent(None, key_code, False)
        Quartz.CGEventPost(Quartz.kCGHIDEventTap, up)
        if shift:
            shift_event = Quartz.CGEventCreateKeyboardEvent(None, 56, False)
            Quartz.CGEventPost(Quartz.kCGHIDEventTap, shift_event)
    return {"success": True, "text": text}


def keyboard_shortcut(modifiers, key):
    """Press keyboard shortcut. modifiers: list of 'cmd','shift','ctrl','alt'."""
    import Quartz

    mod_map = {"cmd": 55, "shift": 56, "ctrl": 59, "alt": 58}
    key_map = {
        "c": 8,
        "v": 9,
        "x": 7,
        "z": 6,
        "a": 0,
        "s": 1,
        "f": 3,
        "p": 35,
        "n": 45,
        "w": 13,
        "q": 12,
        "r": 15,
        "t": 17,
        "enter": 36,
        "tab": 48,
        "esc": 53,
        "space": 49,
        "delete": 51,
    }
    key_code = key_map.get(key.lower(), 0)
    if key_code == 0:
        return {"success": False, "error": f"Unknown key: {key}"}
    mod_codes = [mod_map[m.lower()] for m in modifiers if m.lower() in mod_map]
    for mc in mod_codes:
        Quartz.CGEventPost(
            Quartz.kCGHIDEventTap, Quartz.CGEventCreateKeyboardEvent(None, mc, True)
        )
    import time

    time.sleep(0.05)
    Quartz.CGEventPost(
        Quartz.kCGHIDEventTap, Quartz.CGEventCreateKeyboardEvent(None, key_code, True)
    )
    Quartz.CGEventPost(
        Quartz.kCGHIDEventTap, Quartz.CGEventCreateKeyboardEvent(None, key_code, False)
    )
    for mc in reversed(mod_codes):
        Quartz.CGEventPost(
            Quartz.kCGHIDEventTap, Quartz.CGEventCreateKeyboardEvent(None, mc, False)
        )
    return {"success": True, "shortcut": f"{'+'.join(modifiers)}+{key}"}


# ─── Worker Window ─────────────────────────────────────────────────────────


def open_worker_window(title=None):
    """Open a Terminal window with CursorEye Worker label and custom branding."""
    window_title = title or WORKER_WINDOW_TITLE
    scpt = tempfile.NamedTemporaryFile(suffix=".scpt", delete=False, mode="w")
    scpt.write(f'''
on run
	tell application "Terminal"
		activate
		set workerWin to do script "printf '\\\\033]1337;SetColors=bg=#000000,fg=#d4b8ff\\\\007' && clear && echo '' && echo '  ██████╗██╗   ██╗██████╗ ███████╗██████╗ ███████╗██████╗ ██╗   ██╗███╗   ██╗██╗  ██╗' && echo ' ██╔════╝╚██╗ ██╔╝██╔══██╗██╔════╝██╔══██╗██╔════╝██╔══██╗██║   ██║████╗  ██║██║ ██╔╝' && echo ' ██║      ╚████╔╝ ██████╔╝█████╗  ██████╔╝█████╗  ██║  ██║██║   ██║██╔██╗ ██║█████╔╝ ' && echo ' ██║       ╚██╔╝  ██╔══██╗██╔══╝  ██╔══██╗██╔══╝  ██║  ██║██║   ██║██║╚██╗██║██╔═██╗ ' && echo ' ╚██████╗   ██║   ██████╔╝███████╗██║  ██║███████╗██████╔╝╚██████╔╝██║ ╚████║██║  ██╗' && echo '  ╚═════╝   ╚═╝   ╚═════╝ ╚══════╝╚═╝  ╚═╝╚══════╝╚═════╝  ╚═════╝ ╚═╝  ╚═══╝╚═╝  ╚═╝' && echo '' && echo '  ────────────────────────────────────────' && echo '  AI Workspace • Autonomous Terminal' && echo '  ────────────────────────────────────────' && echo '' && bash"
		set custom title of front window to "{window_title}"
	end tell
end run
''')
    scpt.close()
    result = subprocess.run(["osascript", scpt.name], capture_output=True, text=True)
    os.unlink(scpt.name)
    import time

    time.sleep(1.0)
    return {"success": result.returncode == 0, "window_title": window_title}


def type_in_worker_window(text):
    """Type text directly into the Worker window via AppleScript keystroke.
    This uses Accessibility API to target the specific window, so it does NOT
    interfere with the user's keyboard input in other apps (e.g. games)."""
    escaped = text.replace("\\", "\\\\").replace('"', '\\"').replace("'", "'\\''")
    script = f'''
    tell application "System Events"
        set p to first process whose name is "Terminal"
        set w to first window of p whose name contains "CursorEye Worker"
        set focused of w to true
        set index of w to 1
        delay 0.1
        keystroke "{escaped}"
    end tell
    '''
    result = subprocess.run(["osascript", "-e", script], capture_output=True, text=True)
    if result.returncode != 0:
        return {
            "success": False,
            "error": f"AppleScript keystroke failed: {result.stderr[:200]}",
        }
    return {"success": True, "text": text}


def run_in_worker_window(command):
    """Run a shell command in the worker window."""
    scpt = tempfile.NamedTemporaryFile(suffix=".scpt", delete=False, mode="w")
    scpt.write(f'''
on run
	tell application "Terminal"
		set index of every window whose custom title contains "CursorEye Worker" to 1
		do script "{command}" in first tab of front window
	end tell
end run
''')
    scpt.close()
    result = subprocess.run(["osascript", scpt.name], capture_output=True, text=True)
    os.unlink(scpt.name)
    return {"success": result.returncode == 0}


def shortcut_in_worker_window(modifiers, key):
    """Send keyboard shortcut directly to Worker window via AppleScript.
    Uses Accessibility API — does NOT interfere with user's keyboard in other apps."""
    mod_flags = []
    for m in modifiers:
        ml = m.lower()
        if ml == "cmd":
            mod_flags.append("command down")
        elif ml == "shift":
            mod_flags.append("shift down")
        elif ml == "ctrl":
            mod_flags.append("control down")
        elif ml == "alt":
            mod_flags.append("option down")
    using_clause = ", ".join(mod_flags) if mod_flags else ""
    using_str = f" using {{" + using_clause + "}}" if using_clause else ""
    escaped_key = key.replace("\\", "\\\\").replace('"', '\\"')
    script = f'''
    tell application "System Events"
        set p to first process whose name is "Terminal"
        set w to first window of p whose name contains "CursorEye Worker"
        set focused of w to true
        set index of w to 1
        delay 0.1
        keystroke "{escaped_key}"{using_str}
    end tell
    '''
    result = subprocess.run(["osascript", "-e", script], capture_output=True, text=True)
    if result.returncode != 0:
        return {
            "success": False,
            "error": f"AppleScript shortcut failed: {result.stderr[:200]}",
        }
    return {"success": True, "shortcut": f"{'+'.join(modifiers)}+{key}"}


def run_in_worker_and_wait(command, wait_seconds=3):
    """Run a shell command in the worker window, wait, then OCR the output."""
    run_result = run_in_worker_window(command)
    if not run_result["success"]:
        return {"success": False, "error": "Failed to run command in worker window"}
    import time

    time.sleep(wait_seconds)
    ocr_result = ocr_worker()
    return {
        "success": True,
        "action": "worker_run_wait",
        "command": command,
        "output": ocr_result.get("text", ""),
        "ocr_blocks": ocr_result.get("blocks", 0),
    }


# ─── AI Purple Cursor ──────────────────────────────────────────────────────


def _start_purple_cursor_overlay():
    """Launch the purple cursor overlay as a background process."""
    if os.path.exists(PURPLE_CURSOR_PID_FILE):
        try:
            with open(PURPLE_CURSOR_PID_FILE) as f:
                pid = int(f.read().strip())
            os.kill(pid, 0)
            return pid
        except (ProcessLookupError, ValueError, FileNotFoundError):
            pass

    cursor_script = os.path.join(
        os.path.dirname(os.path.abspath(__file__)), "_purple_cursor.py"
    )
    if not os.path.exists(cursor_script):
        return None

    proc = subprocess.Popen(
        ["python3", cursor_script],
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
    )
    with open(PURPLE_CURSOR_PID_FILE, "w") as f:
        f.write(str(proc.pid))
    return proc.pid


def _stop_purple_cursor_overlay():
    """Stop the purple cursor overlay process."""
    if os.path.exists(PURPLE_CURSOR_PID_FILE):
        try:
            with open(PURPLE_CURSOR_PID_FILE) as f:
                pid = int(f.read().strip())
            os.kill(pid, 9)
        except (ProcessLookupError, ValueError, FileNotFoundError):
            pass
        try:
            os.unlink(PURPLE_CURSOR_PID_FILE)
        except OSError:
            pass


def worker_click(x_pct, y_pct, button="left", click_count=1):
    """Click at percentage position within the Worker window using AppleScript.
    Uses Accessibility API click — does NOT move the system cursor.
    Shows purple cursor overlay at click position."""
    bounds = get_worker_window_bounds()
    if not bounds:
        return {
            "success": False,
            "error": "Worker window not found. Open it first with open_worker.",
        }

    abs_x = int(bounds["x"] + (x_pct / 100) * bounds["w"])
    abs_y = int(bounds["y"] + (y_pct / 100) * bounds["h"])

    _save_cursor_pos(x_pct, y_pct)
    _save_worker_bounds(bounds)
    _start_purple_cursor_overlay()

    import time

    time.sleep(0.15)

    btn_map = {"left": "", "right": "right", "center": "center"}
    btn_str = btn_map.get(button, "")
    click_cmd = "click at"
    if click_count == 2:
        click_cmd = "double click at"
    elif click_count == 3:
        click_cmd = "triple click at"

    escaped_btn = f" using {{{btn_str} down}}" if btn_str else ""
    scpt = tempfile.NamedTemporaryFile(suffix=".scpt", delete=False, mode="w")
    scpt.write(f"""
on run
tell application "System Events"
set p to first process whose name is "Terminal"
set w to first window of p whose name contains "CursorEye Worker"
set focused of w to true
set index of w to 1
delay 0.1
{click_cmd} {{{abs_x}, {abs_y}}}
end tell
end run
""")
    scpt.close()
    result = subprocess.run(["osascript", scpt.name], capture_output=True, text=True)
    os.unlink(scpt.name)
    return {
        "success": result.returncode == 0,
        "action": "worker_click",
        "x_pct": x_pct,
        "y_pct": y_pct,
        "abs_x": abs_x,
        "abs_y": abs_y,
        "error": result.stderr[:200] if result.returncode != 0 else None,
    }


def worker_drag(x1_pct, y1_pct, x2_pct, y2_pct, duration=0.5):
    """Drag from one position to another within the Worker window."""
    bounds = get_worker_window_bounds()
    if not bounds:
        return {"success": False, "error": "Worker window not found."}

    abs_x1 = int(bounds["x"] + (x1_pct / 100) * bounds["w"])
    abs_y1 = int(bounds["y"] + (y1_pct / 100) * bounds["h"])
    abs_x2 = int(bounds["x"] + (x2_pct / 100) * bounds["w"])
    abs_y2 = int(bounds["y"] + (y2_pct / 100) * bounds["h"])

    _save_cursor_pos(x2_pct, y2_pct)
    _save_worker_bounds(bounds)
    _start_purple_cursor_overlay()

    scpt = tempfile.NamedTemporaryFile(suffix=".scpt", delete=False, mode="w")
    scpt.write(f"""
on run
tell application "System Events"
set p to first process whose name is "Terminal"
set w to first window of p whose name contains "CursorEye Worker"
set focused of w to true
set index of w to 1
delay 0.1
drag from {{{abs_x1}, {abs_y1}}} to {{{abs_x2}, {abs_y2}}}
end tell
end run
""")
    scpt.close()
    result = subprocess.run(["osascript", scpt.name], capture_output=True, text=True)
    os.unlink(scpt.name)
    return {
        "success": result.returncode == 0,
        "action": "worker_drag",
        "x1_pct": x1_pct,
        "y1_pct": y1_pct,
        "x2_pct": x2_pct,
        "y2_pct": y2_pct,
    }


def worker_scroll(x_pct, y_pct, amount, direction="down"):
    """Scroll within the Worker window at a given position."""
    bounds = get_worker_window_bounds()
    if not bounds:
        return {"success": False, "error": "Worker window not found."}

    abs_x = int(bounds["x"] + (x_pct / 100) * bounds["w"])
    abs_y = int(bounds["y"] + (y_pct / 100) * bounds["h"])

    scroll_cmd = "scroll down" if direction == "down" else "scroll up"

    scpt = tempfile.NamedTemporaryFile(suffix=".scpt", delete=False, mode="w")
    scpt.write(f"""
on run
tell application "System Events"
set p to first process whose name is "Terminal"
set w to first window of p whose name contains "CursorEye Worker"
set focused of w to true
set index of w to 1
delay 0.1
{scroll_cmd} {amount} at {{{abs_x}, {abs_y}}}
end tell
end run
""")
    scpt.close()
    result = subprocess.run(["osascript", scpt.name], capture_output=True, text=True)
    os.unlink(scpt.name)
    return {
        "success": result.returncode == 0,
        "action": "worker_scroll",
        "amount": amount,
        "direction": direction,
    }


def worker_move_cursor(x_pct, y_pct):
    """Move the purple cursor overlay to a position within the Worker window (no click)."""
    bounds = get_worker_window_bounds()
    if not bounds:
        return {"success": False, "error": "Worker window not found."}

    _save_cursor_pos(x_pct, y_pct)
    _save_worker_bounds(bounds)
    _start_purple_cursor_overlay()
    return {
        "success": True,
        "action": "worker_move_cursor",
        "x_pct": x_pct,
        "y_pct": y_pct,
    }


def worker_hide_cursor():
    """Hide the purple cursor overlay."""
    _stop_purple_cursor_overlay()
    return {"success": True, "action": "worker_hide_cursor"}


def _save_cursor_pos(x_pct, y_pct):
    """Save cursor position to shared file for overlay process."""
    os.makedirs(os.path.dirname(PURPLE_CURSOR_POS_FILE), exist_ok=True)
    with open(PURPLE_CURSOR_POS_FILE, "w") as f:
        json.dump({"x_pct": x_pct, "y_pct": y_pct}, f)


def _save_worker_bounds(bounds):
    """Save worker window bounds to shared file for overlay process."""
    bounds_file = os.path.expanduser("~/.cursoreye/worker_bounds.json")
    os.makedirs(os.path.dirname(bounds_file), exist_ok=True)
    with open(bounds_file, "w") as f:
        json.dump(bounds, f)


# ─── File Generation ────────────────────────────────────────────────────────


def generate_file(file_type, content, filename=None):
    """Generate a file and return its path. Supports txt, docx, pdf, md."""
    if not filename:
        ext_map = {
            "txt": "txt",
            "docx": "docx",
            "pdf": "pdf",
            "md": "md",
            "html": "html",
            "csv": "csv",
        }
        filename = f"cursoreye_output.{ext_map.get(file_type, 'txt')}"
    filepath = os.path.join(os.path.expanduser("~/Desktop"), filename)

    if file_type == "txt" or file_type == "md":
        with open(filepath, "w") as f:
            f.write(content)
        return {"success": True, "path": filepath}

    elif file_type == "csv":
        with open(filepath, "w") as f:
            f.write(content)
        return {"success": True, "path": filepath}

    elif file_type == "html":
        html = f"""<!DOCTYPE html>
<html><head><meta charset="utf-8"><title>CursorEye Document</title>
<style>body{{font-family:sans-serif;max-width:800px;margin:40px auto;padding:0 20px;line-height:1.6}}</style>
</head><body>{content}</body></html>"""
        with open(filepath, "w") as f:
            f.write(html)
        return {"success": True, "path": filepath}

    elif file_type == "docx":
        try:
            subprocess.run(
                ["pip3", "install", "python-docx", "-q"], check=True, timeout=30
            )
            from docx import Document

            doc = Document()
            for paragraph in content.split("\n"):
                if paragraph.startswith("# "):
                    doc.add_heading(paragraph[2:], level=1)
                elif paragraph.startswith("## "):
                    doc.add_heading(paragraph[2:], level=2)
                elif paragraph.startswith("### "):
                    doc.add_heading(paragraph[2:], level=3)
                elif paragraph.strip() == "":
                    continue
                else:
                    doc.add_paragraph(paragraph)
            doc.save(filepath)
            return {"success": True, "path": filepath}
        except Exception as e:
            with open(filepath + ".txt", "w") as f:
                f.write(content)
            return {
                "success": True,
                "path": filepath + ".txt",
                "fallback": True,
                "error": str(e),
            }

    elif file_type == "pdf":
        try:
            subprocess.run(["pip3", "install", "fpdf2", "-q"], check=True, timeout=30)
            from fpdf import FPDF

            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Helvetica", size=12)
            for line in content.split("\n"):
                pdf.cell(0, 10, line, new_x="LMARGIN", new_y="NEXT")
            pdf.output(filepath)
            return {"success": True, "path": filepath}
        except Exception as e:
            with open(filepath + ".txt", "w") as f:
                f.write(content)
            return {
                "success": True,
                "path": filepath + ".txt",
                "fallback": True,
                "error": str(e),
            }

    return {"success": False, "error": f"Unsupported file type: {file_type}"}


# ─── Get Active App ────────────────────────────────────────────────────────


def get_active_app():
    """Get currently active application name."""
    script = 'tell application "System Events" to get name of first process whose frontmost is true'
    result = subprocess.run(["osascript", "-e", script], capture_output=True, text=True)
    return result.stdout.strip()


# ─── Command Dispatcher ────────────────────────────────────────────────────


def execute_command(cmd):
    """Execute a CursorEye command and return result."""
    action = cmd.get("action")
    params = cmd.get("params", {})

    try:
        if action == "screenshot":
            img = take_screenshot(params.get("region"))
            return {
                "success": True,
                "action": "screenshot",
                "image": img[:100] + "...",
                "image_length": len(img),
            }

        elif action == "ocr":
            result = ocr_screen(params.get("region"))
            texts = [b["text"] for b in result.get("text_blocks", [])]
            return {
                "success": True,
                "action": "ocr",
                "text": "\n".join(texts),
                "blocks": len(texts),
            }

        elif action == "ocr_full":
            result = ocr_screen(params.get("region"))
            return {
                "success": True,
                "action": "ocr_full",
                "text_blocks": result["text_blocks"],
                "image_length": len(result.get("image", "")),
            }

        elif action == "click":
            return {
                **mouse_click(
                    params.get("x", 50),
                    params.get("y", 50),
                    params.get("button", "left"),
                    params.get("count", 1),
                ),
                "action": "click",
            }

        elif action == "double_click":
            return {
                **mouse_click(params.get("x", 50), params.get("y", 50), "left", 2),
                "action": "double_click",
            }

        elif action == "right_click":
            return {
                **mouse_click(params.get("x", 50), params.get("y", 50), "right", 1),
                "action": "right_click",
            }

        elif action == "drag":
            return {
                **mouse_drag(
                    params.get("x1", 0),
                    params.get("y1", 0),
                    params.get("x2", 0),
                    params.get("y2", 0),
                    params.get("duration", 0.5),
                ),
                "action": "drag",
            }

        elif action == "scroll":
            return {
                **mouse_scroll(
                    params.get("x", 50),
                    params.get("y", 50),
                    params.get("amount", 3),
                    params.get("direction", "down"),
                ),
                "action": "scroll",
            }

        elif action == "type":
            return {**keyboard_type(params.get("text", "")), "action": "type"}

        elif action == "shortcut":
            return {
                **keyboard_shortcut(params.get("modifiers", []), params.get("key", "")),
                "action": "shortcut",
            }

        elif action == "open_worker":
            return {**open_worker_window(params.get("title")), "action": "open_worker"}

        elif action == "worker_type":
            return {
                **type_in_worker_window(params.get("text", "")),
                "action": "worker_type",
            }

        elif action == "worker_shortcut":
            return {
                **shortcut_in_worker_window(
                    params.get("modifiers", []), params.get("key", "")
                ),
                "action": "worker_shortcut",
            }

        elif action == "worker_run":
            return {
                **run_in_worker_window(params.get("command", "")),
                "action": "worker_run",
            }

        elif action == "worker_run_wait":
            return run_in_worker_and_wait(
                params.get("command", ""),
                params.get("wait_seconds", 3),
            )

        elif action == "screenshot_worker":
            return screenshot_worker()

        elif action == "ocr_worker":
            return ocr_worker()

        elif action == "worker_click":
            return worker_click(
                params.get("x", 50),
                params.get("y", 50),
                params.get("button", "left"),
                params.get("count", 1),
            )

        elif action == "worker_double_click":
            return worker_click(
                params.get("x", 50),
                params.get("y", 50),
                "left",
                2,
            )

        elif action == "worker_right_click":
            return worker_click(
                params.get("x", 50),
                params.get("y", 50),
                "right",
                1,
            )

        elif action == "worker_drag":
            return worker_drag(
                params.get("x1", 0),
                params.get("y1", 0),
                params.get("x2", 0),
                params.get("y2", 0),
                params.get("duration", 0.5),
            )

        elif action == "worker_scroll":
            return worker_scroll(
                params.get("x", 50),
                params.get("y", 50),
                params.get("amount", 3),
                params.get("direction", "down"),
            )

        elif action == "worker_move_cursor":
            return worker_move_cursor(
                params.get("x", 50),
                params.get("y", 50),
            )

        elif action == "worker_hide_cursor":
            return worker_hide_cursor()

        elif action == "generate_file":
            return {
                **generate_file(
                    params.get("file_type", "txt"),
                    params.get("content", ""),
                    params.get("filename"),
                ),
                "action": "generate_file",
            }

        elif action == "get_active_app":
            app = get_active_app()
            return {"success": True, "action": "get_active_app", "app": app}

        elif action == "verify":
            result = ocr_screen(params.get("region"))
            texts = [b["text"] for b in result.get("text_blocks", [])]
            full_text = "\n".join(texts)
            target = params.get("target", "").lower()
            found = target in full_text.lower() if target else True
            return {
                "success": True,
                "action": "verify",
                "found": found,
                "screen_text": full_text[:500],
                "target": target,
            }

        elif action == "open_app":
            app_name = params.get("app", "")
            result = subprocess.run(
                ["open", "-a", app_name], capture_output=True, text=True
            )
            return {
                "success": result.returncode == 0,
                "action": "open_app",
                "app": app_name,
            }

        else:
            return {"success": False, "error": f"Unknown action: {action}"}

    except Exception as e:
        return {"success": False, "action": action, "error": str(e)}


# ─── WebSocket Server ──────────────────────────────────────────────────────


async def handle_client(websocket):
    print(f"[+] Client connected")
    try:
        async for message in websocket:
            try:
                cmd = json.loads(message)
                seq = cmd.pop("_seq", None)
                result = execute_command(cmd)
                if seq is not None:
                    result["_seq"] = seq
                await websocket.send(json.dumps(result))
            except json.JSONDecodeError:
                await websocket.send(
                    json.dumps({"success": False, "error": "Invalid JSON"})
                )
    except websockets.exceptions.ConnectionClosed:
        print(f"[-] Client disconnected")


async def main():
    ws_url = sys.argv[1] if len(sys.argv) > 1 else DEFAULT_WS_URL
    from urllib.parse import urlparse

    parsed = urlparse(ws_url)
    host = parsed.hostname or "localhost"
    port = parsed.port or 8765

    print(f"🔴 CursorEye Agent v{AGENT_VERSION}")
    print(f" Listening on ws://{host}:{port}")
    print(f" Capabilities: screenshot, ocr, click, double_click, right_click,")
    print(f" drag, scroll, type, shortcut, open_worker, worker_type,")
    print(f" worker_run, worker_run_wait, screenshot_worker, ocr_worker,")
    print(f" worker_click, worker_double_click, worker_right_click,")
    print(f" worker_drag, worker_scroll, worker_move_cursor, worker_hide_cursor,")
    print(f" generate_file, get_active_app, verify, open_app")
    print(f"   Press Ctrl+C to stop")
    print()

    async with websockets.serve(handle_client, host, port):
        await asyncio.Future()


if __name__ == "__main__":
    asyncio.run(main())
